import os
from datetime import datetime
from datetime import timezone
from pathlib import Path
from typing import Any
from typing import IO

from onyx.configs.app_configs import INDEX_BATCH_SIZE
from onyx.configs.constants import DocumentSource
from onyx.configs.constants import FileOrigin
from onyx.connectors.cross_connector_utils.miscellaneous_utils import (
    process_onyx_metadata,
)
from onyx.connectors.interfaces import GenerateDocumentsOutput
from onyx.connectors.interfaces import LoadConnector
from onyx.connectors.models import Document
from onyx.connectors.models import ImageSection
from onyx.connectors.models import TextSection
from onyx.file_processing.extract_file_text import extract_text_and_images
from onyx.file_processing.extract_file_text import get_file_ext
from onyx.file_processing.extract_file_text import is_accepted_file_ext
from onyx.file_processing.extract_file_text import OnyxExtensionType
from onyx.file_processing.image_utils import store_image_and_create_section
from onyx.file_store.file_store import get_default_file_store
from onyx.utils.logger import setup_logger

##################################################################################################################################

from PyPDF2 import PdfReader
from pptx import Presentation
from openpyx1 import load_workbook
from docx import Document
from onyx.db.models import IndexAttempt
from onyx.db import get_db
from pathlib import Path
from onyx.connectors.interfaces import GenerateDocumentsOutput, LoadConnector
from onyx.file_processing.ollama_ocr import get_ollama_ocr

############################################################################################################################################

logger = setup_logger()

def get_file_metadata(file_path, file_type):
    if file_type == 'pdf':
        try:
            with open(file_path, 'rb') as file:
                pdf = PdfReader(file)
                return {'total_pages': len(pdf.pages)}
        except Exception as e:
            logger.error(f"Error extracting PDF metadata: {e}")
            return {'total_pages': 0}
    elif file_type == 'pptx':
        try:
            prs = Presentation(file_path)
            return {'total_slides': len(prs.slides)}
        except Exception as e:
            logger.error(f"Error extracting PPTX metadata: {e}")
            return {'total_slides': 0}
    elif file_type == 'xlsx':
        try:
            wb = load_workbook(file_path)
            return {'total_sheets': len(wb.sheetnames)}
        except Exception as e:
            logger.error(f"Error extracting XLSX metadata: {e}")
            return {'total_sheets': 0}
    elif file_type == 'docx':
        try:
            doc = Document(file_path)
            return {'total_sections': 1}  # Default to single batch
        except Exception as e:
            logger.error(f"Error extracting DOCX metadata: {e}")
            return {'total_sections': 0}
    return {}

def _create_image_section(
    image_data: bytes,
    parent_file_name: str,
    display_name: str,
    link: str | None = None,
    idx: int = 0,
) -> tuple[ImageSection, str | None]:
    """
    Creates an ImageSection for an image file or embedded image.
    Stores the image in FileStore but does not generate a summary.

    Args:
        image_data: Raw image bytes
        db_session: Database session
        parent_file_name: Name of the parent file (for embedded images)
        display_name: Display name for the image
        idx: Index for embedded images

    Returns:
        Tuple of (ImageSection, stored_file_name or None)
    """
    # Create a unique identifier for the image
    file_id = f"{parent_file_name}_embedded_{idx}" if idx > 0 else parent_file_name

    # Store the image and create a section
    try:
        section, stored_file_name = store_image_and_create_section(
            image_data=image_data,
            file_id=file_id,
            display_name=display_name,
            link=link,
            file_origin=FileOrigin.CONNECTOR,
        )
        return section, stored_file_name
    except Exception as e:
        logger.error(f"Failed to store image {display_name}: {e}")
        raise e


# def _process_file(
#     file_id: str,
#     file_name: str,
#     file: IO[Any],
#     metadata: dict[str, Any] | None,
#     pdf_pass: str | None,
# ) -> list[Document]:
#     """
#     Process a file and return a list of Documents.
#     For images, creates ImageSection objects without summarization.
#     For documents with embedded images, extracts and stores the images.
#     """
#     if metadata is None:
#         metadata = {}

#     # Get file extension and determine file type
#     extension = get_file_ext(file_name)

#     if not is_accepted_file_ext(extension, OnyxExtensionType.All):
#         logger.warning(
#             f"Skipping file '{file_name}' with unrecognized extension '{extension}'"
#         )
#         return []

#     # If a zip is uploaded with a metadata file, we can process it here
#     onyx_metadata, custom_tags = process_onyx_metadata(metadata)
#     file_display_name = onyx_metadata.file_display_name or os.path.basename(file_name)
#     time_updated = onyx_metadata.doc_updated_at or datetime.now(timezone.utc)
#     primary_owners = onyx_metadata.primary_owners
#     secondary_owners = onyx_metadata.secondary_owners
#     link = onyx_metadata.link

#     # These metadata items are not settable by the user
#     source_type_str = metadata.get("connector_type")
#     source_type = (
#         DocumentSource(source_type_str) if source_type_str else DocumentSource.FILE
#     )

#     doc_id = f"FILE_CONNECTOR__{file_id}"
#     title = metadata.get("title") or file_display_name

#     # 1) If the file itself is an image, handle that scenario quickly
#     if extension in LoadConnector.IMAGE_EXTENSIONS:
#         # Read the image data
#         image_data = file.read()
#         if not image_data:
#             logger.warning(f"Empty image file: {file_name}")
#             return []

#         # Create an ImageSection for the image
#         try:
#             section, _ = _create_image_section(
#                 image_data=image_data,
#                 parent_file_name=file_id,
#                 display_name=title,
#             )

#             return [
#                 Document(
#                     id=doc_id,
#                     sections=[section],
#                     source=source_type,
#                     semantic_identifier=file_display_name,
#                     title=title,
#                     doc_updated_at=time_updated,
#                     primary_owners=primary_owners,
#                     secondary_owners=secondary_owners,
#                     metadata=custom_tags,
#                 )
#             ]
#         except Exception as e:
#             logger.error(f"Failed to process image file {file_name}: {e}")
#             return []

#     # 2) Otherwise: text-based approach. Possibly with embedded images.
#     file.seek(0)

#     # Extract text and images from the file
#     extraction_result = extract_text_and_images(
#         file=file,
#         file_name=file_name,
#         pdf_pass=pdf_pass,
#     )

#     # Each file may have file-specific ONYX_METADATA https://docs.onyx.app/connectors/file
#     # If so, we should add it to any metadata processed so far
#     if extraction_result.metadata:
#         logger.debug(
#             f"Found file-specific metadata for {file_name}: {extraction_result.metadata}"
#         )
#         onyx_metadata, more_custom_tags = process_onyx_metadata(
#             extraction_result.metadata
#         )

#         # Add file-specific tags
#         custom_tags.update(more_custom_tags)

#         # File-specific metadata overrides metadata processed so far
#         source_type = onyx_metadata.source_type or source_type
#         primary_owners = onyx_metadata.primary_owners or primary_owners
#         secondary_owners = onyx_metadata.secondary_owners or secondary_owners
#         time_updated = onyx_metadata.doc_updated_at or time_updated
#         file_display_name = onyx_metadata.file_display_name or file_display_name
#         title = onyx_metadata.title or onyx_metadata.file_display_name or title
#         link = onyx_metadata.link or link

#     # Build sections: first the text as a single Section
#     sections: list[TextSection | ImageSection] = []
#     if extraction_result.text_content.strip():
#         logger.debug(f"Creating TextSection for {file_name} with link: {link}")
#         sections.append(
#             TextSection(link=link, text=extraction_result.text_content.strip())
#         )

#     # Then any extracted images from docx, PDFs, etc.
#     for idx, (img_data, img_name) in enumerate(
#         extraction_result.embedded_images, start=1
#     ):
#         # Store each embedded image as a separate file in FileStore
#         # and create a section with the image reference
#         try:
#             image_section, stored_file_name = _create_image_section(
#                 image_data=img_data,
#                 parent_file_name=file_id,
#                 display_name=f"{title} - image {idx}",
#                 idx=idx,
#             )
#             sections.append(image_section)
#             logger.debug(
#                 f"Created ImageSection for embedded image {idx} "
#                 f"in {file_name}, stored as: {stored_file_name}"
#             )
#         except Exception as e:
#             logger.warning(
#                 f"Failed to process embedded image {idx} in {file_name}: {e}"
#             )

#     return [
#         Document(
#             id=doc_id,
#             sections=sections,
#             source=source_type,
#             semantic_identifier=file_display_name,
#             title=title,
#             doc_updated_at=time_updated,
#             primary_owners=primary_owners,
#             secondary_owners=secondary_owners,
#             metadata=custom_tags,
#         )
#     ]
#################################################################################################################################################################################################################

def _process_file(file_id: str, file_name: str, file: IO[Any], metadata: dict[str, Any] | None, pdf_pass: str | None, index_attempt_id: int) -> list[Document]:
    if metadata is None:
        metadata = {}
    extension = get_file_ext(file_name)
    if not is_accepted_file_ext(extension, OnyxExtensionType.All):
        logger.warning(f"Skipping file '{file_name}' with unrecognized extension '{extension}'")
        return []
    onyx_metadata, custom_tags = process_onyx_metadata(metadata)
    file_display_name = onyx_metadata.file_display_name or os.path.basename(file_name)
    time_updated = onyx_metadata.doc_updated_at or datetime.now(timezone.utc)
    primary_owners = onyx_metadata.primary_owners
    secondary_owners = onyx_metadata.secondary_owners
    link = onyx_metadata.link
    source_type = DocumentSource(metadata.get("connector_type") or "FILE")
    doc_id = f"FILE_CONNECTOR__{file_id}"
    title = metadata.get("title") or file_display_name

    if extension in LoadConnector.IMAGE_EXTENSIONS:
        image_data = file.read()
        if not image_data:
            logger.warning(f"Empty image file: {file_name}")
            return []
        section, _ = _create_image_section(image_data, file_id, title)
        with get_db() as db:
            index_attempt = db.query(IndexAttempt).filter(IndexAttempt.id == index_attempt_id).first()
            if index_attempt:
                index_attempt.file_type = extension
                index_attempt.completed_batches += 1
                index_attempt.is_complete = True
                db.commit()
        return [Document(id=doc_id, sections=[section], source=source_type, semantic_identifier=file_display_name, title=title, doc_updated_at=time_updated, primary_owners=primary_owners, secondary_owners=secondary_owners, metadata=custom_tags)]

    file.seek(0)
    file_path = f"/tmp/{file_id}{os.path.splitext(file_name)[1]}"
    with open(file_path, "wb") as f:
        f.write(file.read())
    file.seek(0)  # Rewind for reuse

    if extension == 'pdf':
        ocr = get_ollama_ocr()
        if ocr and ocr.available:
            text_content = ocr.extract_text_from_pdf(file)
            sections = [TextSection(link=link, text=text_content.strip())] if text_content.strip() else []
            with get_db() as db:
                index_attempt = db.query(IndexAttempt).filter(IndexAttempt.id == index_attempt_id).first()
                if index_attempt:
                    index_attempt.file_type = 'pdf'
                    index_attempt.total_pages = min(ocr.max_pages_per_pdf, len(fitz.open(stream=file.read(), filetype="pdf")))  # Total pages
                    index_attempt.processed_pages = index_attempt.total_pages  # All pages processed by OCR
                    index_attempt.is_complete = True
                    db.commit()
            return [Document(id=doc_id, sections=sections, source=source_type, semantic_identifier=file_display_name, title=title, doc_updated_at=time_updated, primary_owners=primary_owners, secondary_owners=secondary_owners, metadata=custom_tags)]
        else:
            logger.warning(f"Ollama OCR not available for {file_name}, falling back to extract_text_and_images")
            extraction_result = extract_text_and_images(file=file, file_name=file_name, pdf_pass=pdf_pass)
    else:
        extraction_result = extract_text_and_images(file=file, file_name=file_name, pdf_pass=pdf_pass)

    if extraction_result.metadata:
        onyx_metadata, more_custom_tags = process_onyx_metadata(extraction_result.metadata)
        custom_tags.update(more_custom_tags)
        source_type = onyx_metadata.source_type or source_type
        primary_owners = onyx_metadata.primary_owners or primary_owners
        secondary_owners = onyx_metadata.secondary_owners or secondary_owners
        time_updated = onyx_metadata.doc_updated_at or time_updated
        file_display_name = onyx_metadata.file_display_name or file_display_name
        title = onyx_metadata.title or onyx_metadata.file_display_name or title
        link = onyx_metadata.link or link

    sections = []
    if extraction_result.text_content.strip():
        sections.append(TextSection(link=link, text=extraction_result.text_content.strip()))

    for idx, (img_data, img_name) in enumerate(extraction_result.embedded_images, start=1):
        try:
            image_section, stored_file_name = _create_image_section(img_data, file_id, f"{title} - image {idx}", idx=idx)
            sections.append(image_section)
            logger.debug(f"Created ImageSection for embedded image {idx} in {file_name}, stored as: {stored_file_name}")
        except Exception as e:
            logger.warning(f"Failed to process embedded image {idx} in {file_name}: {e}")

    with get_db() as db:
        index_attempt = db.query(IndexAttempt).filter(IndexAttempt.id == index_attempt_id).first()
        if index_attempt:
            file_type = extension
            index_attempt.file_type = file_type
            metadata_result = get_file_metadata(file_path, file_type)
            index_attempt.total_pages = metadata_result.get('total_pages', 0)
            index_attempt.total_slides = metadata_result.get('total_slides', 0)
            index_attempt.total_sheets = metadata_result.get('total_sheets', 0)
            index_attempt.total_sections = metadata_result.get('total_sections', 0)
            if file_type == 'pdf' and not ocr:  # Fallback case
                index_attempt.processed_pages = index_attempt.total_pages
            elif file_type == 'pptx' and extraction_result.total_slides:
                index_attempt.processed_slides = extraction_result.total_slides
            elif file_type == 'xlsx' and extraction_result.total_sheets:
                index_attempt.processed_sheets = extraction_result.total_sheets
            elif file_type == 'docx':
                index_attempt.processed_sections = 1
            else:
                index_attempt.completed_batches += 1
            if (index_attempt.processed_pages == index_attempt.total_pages or
                index_attempt.processed_slides == index_attempt.total_slides or
                index_attempt.processed_sheets == index_attempt.total_sheets or
                index_attempt.processed_sections == index_attempt.total_sections or
                index_attempt.completed_batches == index_attempt.total_batches):
                index_attempt.is_complete = True
            db.commit()

    return [
        Document(
            id=doc_id,
            sections=sections,
            source=source_type,
            semantic_identifier=file_display_name,
            title=title,
            doc_updated_at=time_updated,
            primary_owners=primary_owners,
            secondary_owners=secondary_owners,
            metadata=custom_tags,
        )
    ]
#########################################################################################################################################################################################################################################################################

# class LocalFileConnector(LoadConnector):
#     """
#     Connector that reads files from Postgres and yields Documents, including
#     embedded image extraction without summarization.
#     """

#     def __init__(
#         self,
#         file_locations: list[Path | str],
#         zip_metadata: dict[str, Any],
#         batch_size: int = INDEX_BATCH_SIZE,
#     ) -> None:
#         self.file_locations = [str(loc) for loc in file_locations]
#         self.batch_size = batch_size
#         self.pdf_pass: str | None = None
#         self.zip_metadata = zip_metadata

#     def load_credentials(self, credentials: dict[str, Any]) -> dict[str, Any] | None:
#         self.pdf_pass = credentials.get("pdf_password")

#         return None

#     def _get_file_metadata(self, file_name: str) -> dict[str, Any]:
#         return self.zip_metadata.get(file_name, {}) or self.zip_metadata.get(
#             os.path.basename(file_name), {}
#         )

#     def load_from_state(self) -> GenerateDocumentsOutput:
#         """
#         Iterates over each file path, fetches from Postgres, tries to parse text
#         or images, and yields Document batches.
#         """
#         documents: list[Document] = []

#         for file_id in self.file_locations:
#             file_store = get_default_file_store()
#             file_record = file_store.read_file_record(file_id=file_id)
#             if not file_record:
#                 # typically an unsupported extension
#                 logger.warning(f"No file record found for '{file_id}' in PG; skipping.")
#                 continue

#             metadata = self._get_file_metadata(file_record.display_name)
#             file_io = file_store.read_file(file_id=file_id, mode="b")
#             new_docs = _process_file(
#                 file_id=file_id,
#                 file_name=file_record.display_name,
#                 file=file_io,
#                 metadata=metadata,
#                 pdf_pass=self.pdf_pass,
#             )
#             documents.extend(new_docs)

#             if len(documents) >= self.batch_size:
#                 yield documents

#                 documents = []

#         if documents:
#             yield documents
                            ################################################################################################################################################################################################################

class LocalFileConnector(LoadConnector):
    def __init__(self, file_locations: list[Path | str], zip_metadata: dict[str, Any], batch_size: int = INDEX_BATCH_SIZE, index_attempt_id: int = None):
        self.file_locations = [str(loc) for loc in file_locations]
        self.batch_size = batch_size
        self.pdf_pass: str | None = None
        self.zip_metadata = zip_metadata
        self.index_attempt_id = index_attempt_id

    def load_credentials(self, credentials: dict[str, Any]) -> dict[str, Any] | None:
        self.pdf_pass = credentials.get("pdf_password")
        return None

    def _get_file_metadata(self, file_name: str) -> dict[str, Any]:
        return self.zip_metadata.get(file_name, {}) or self.zip_metadata.get(os.path.basename(file_name), {})

    def load_from_state(self) -> GenerateDocumentsOutput:
        documents: list[Document] = []
        file_store = get_default_file_store()

        for file_id in self.file_locations:
            file_record = file_store.read_file_record(file_id=file_id)
            if not file_record:
                logger.warning(f"No file record found for '{file_id}' in PG; skipping.")
                continue
            metadata = self._get_file_metadata(file_record.display_name)
            file_io = file_store.read_file(file_id=file_id, mode="b")
            new_docs = _process_file(file_id, file_record.display_name, file_io, metadata, self.pdf_pass, self.index_attempt_id or 0)
            documents.extend(new_docs)
            if len(documents) >= self.batch_size:
                yield documents
                documents = []
        if documents:
            yield documents
#################################################################################################################


if __name__ == "__main__":
    connector = LocalFileConnector(
        file_locations=[os.environ["TEST_FILE"]], zip_metadata={}, index_attempt_id=1        #### added index_attemp_id
    )
    connector.load_credentials({"pdf_password": os.environ.get("PDF_PASSWORD")})
    doc_batches = connector.load_from_state()
    for batch in doc_batches:
        print("BATCH:", batch)
